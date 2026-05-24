from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from jinja2 import Template
import os
from io import BytesIO
from datetime import datetime
import logging
import re
from ..report_utils.diagrams import (
    comparison_bar_chart_docx,
    concentration_dynamics_lineplot_docx
)

logger = logging.getLogger(__name__)

class DocxGenerator:
    """Генератор DOCX-отчетов"""

    def __init__(self):
        self.BASE_DIR = os.path.dirname(os.path.abspath(__file__))
        # константы для полей
        self.LEFT_MARGIN = Inches(30 / 25.4)          # 30mm 
        self.RIGHT_MARGIN = Inches(15 / 25.4)         # 15mm 
        self.TOP_MARGIN = Inches(20 / 25.4)           # 20mm 
        self.BOTTOM_MARGIN = Inches(20 / 25.4)        # 20mm 
        self.FIRST_LINE_INDENT = Inches(12.5 / 25.4)  # 12.5mm

    def format_date_ddmmyyyy(self, value):
        """Преобразование даты из формата YYYY-MM-DD в DD.MM.YYYY"""
        if not value:
            return ""
        try:
            dt = datetime.strptime(value, "%Y-%m-%d")
            return dt.strftime("%d.%m.%Y")
        except Exception:
            return value

    def format_number(self, value):
        """Форматирование числа с двумя знаками после запятой"""
        if value is None or value == "":
            return ""

        try:
            number = float(str(value).replace(",", "."))
            return f"{number:.2f}"
        except (ValueError, TypeError):
            return str(value)
    
    def get_normative_range(self, indicator_name):
        """Возвращает норматив по названию показателя"""

        norms = {
            "ph": "6.00 - 9.00",
            "железо": "0.27 - 0.33",
            "iron": "0.27 - 0.33",
            "марганец": "0.09 - 0.12",
            "manganese": "0.09 - 0.12",
            "нитраты": "38.25 - 51.75",
            "nitrates": "38.25 - 51.75",
            "сульфаты": "435.00 - 565.00",
            "sulfates": "435.00 - 565.00",
        }

        key = str(indicator_name).strip().lower()
        return norms.get(key, "")

    def remove_graph_desc_prefix(self, text):
        """Удаление префикса GRAPH_DESC:"""
        if not text:
            return text

        return re.sub(r"\s*GRAPH_DESC\s*:?\s*", "", text).strip()
    
    def add_page_number(self, paragraph):
        """Добавление номера страницы в нижний колонтитул"""
        paragraph.alignment = (WD_ALIGN_PARAGRAPH.CENTER)
        run = paragraph.add_run()
        run.font.name = ("Times New Roman")
        run.font.size = Pt(12)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"),"begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"),"preserve")
        instr_text.text = " PAGE "
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"),"end")
        run._r.append(fld_begin)
        run._r.append(instr_text)
        run._r.append(fld_end)

    def set_margins(self, doc: Document):
        """Установка полей документа"""
        sections = doc.sections
        for section in sections:
            section.left_margin = self.LEFT_MARGIN
            section.right_margin = self.RIGHT_MARGIN
            section.top_margin = self.TOP_MARGIN
            section.bottom_margin = self.BOTTOM_MARGIN

    def add_title_page(self, doc: Document, user_data: dict):
        """Добавление титульной страницы"""
        # Заголовок организации 
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_run = header_para.add_run(user_data.get("FULL_OBJECT_NAME", ""))
        header_run.font.name = "Times New Roman"
        header_run.font.size = Pt(14)
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Нижний колонтитул 
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_run = footer_para.add_run(f"ПРЕДПРИЯТИЕ {user_data.get('SHORT_OBJECT_NAME', '')}")
        footer_run.font.name = "Times New Roman"
        footer_run.font.size = Pt(14)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Титульная страница, пространство перед заголовком
        for _ in range(12):
            doc.add_paragraph()

        # Основной заголовок
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("ОТЧЕТ ПО ЭКОЛОГИЧЕСКОЙ БЕЗОПАСНОСТИ")
        title_run.font.name = "Times New Roman"
        title_run.font.size = Pt(14)
        title_run.bold = True

        subtitle1_para = doc.add_paragraph()
        subtitle1_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle1_run = subtitle1_para.add_run("ДРЕНАЖНЫХ СИСТЕМ")
        subtitle1_run.font.name = "Times New Roman"
        subtitle1_run.font.size = Pt(14)
        subtitle1_run.bold = True

        subtitle2_para = doc.add_paragraph()
        subtitle2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle2_run = subtitle2_para.add_run(f"ЗА {user_data.get('YEAR', '')} ГОД")
        subtitle2_run.font.name = "Times New Roman"
        subtitle2_run.font.size = Pt(14)
        subtitle2_run.bold = True

        # Новая секция после титульника
        new_section = doc.add_section(WD_SECTION.NEW_PAGE)
        new_section.left_margin = self.LEFT_MARGIN
        new_section.right_margin = self.RIGHT_MARGIN
        new_section.top_margin = self.TOP_MARGIN
        new_section.bottom_margin = self.BOTTOM_MARGIN

        # Убираем header/footer linkage
        new_section.different_first_page_header_footer = False

        # Убираем верхний колонтитул
        new_section.header.is_linked_to_previous = False
        header_para = new_section.header.paragraphs[0]
        header_para.text = ""

        # Нижний колонтитул — только номер страницы
        new_section.footer.is_linked_to_previous = False
        footer_para = new_section.footer.paragraphs[0]
        footer_para.text = ""

        self.add_page_number(footer_para)

    def filter_dynamic_columns(self, data_list, fields):
        """Оставляет только те колонки, где есть данные"""
        used_fields = []
        for f in fields:
            if any(self.has_value(item.get(f)) for item in data_list):
                used_fields.append(f)
        return used_fields

    def has_value(self, value):
        return value is not None and str(value).strip() != ""

    def style_table(self, table):
        """Единый стиль таблиц"""
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                tc_pr = cell._element.get_or_add_tcPr()
                
                # Удаляем старые границы
                existing_borders = tc_pr.find(".//w:tcBorders",tc_pr.nsmap)
                if existing_borders is not None:
                    tc_pr.remove(existing_borders)
                
                # Создаём новые границы 
                borders = OxmlElement("w:tcBorders")
                for border_name in ["top", "left", "bottom", "right"]:
                    border = OxmlElement(f"w:{border_name}")
                    border.set(qn("w:val"), "single")
                    border.set(qn("w:sz"), "8")
                    border.set(qn("w:space"), "0")
                    border.set(qn("w:color"), "000000")
                    borders.append(border)
                tc_pr.append(borders)
                
                # Удаляем старую заливку  
                existing_shd = tc_pr.find(".//w:shd",tc_pr.nsmap)
                if existing_shd is not None:
                    tc_pr.remove(existing_shd)
                
                # Новая заливка  
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                if row_idx == 0:
                    shd.set(qn("w:fill"), "E6E6E6")
                else:
                    shd.set(qn("w:fill"), "FFFFFF")
                tc_pr.append(shd)

                # Формат текста  
                for paragraph in cell.paragraphs:
                    paragraph.alignment = (WD_ALIGN_PARAGRAPH.CENTER)
                    for run in paragraph.runs:
                        run.font.name = ("Times New Roman")
                        run.font.size = Pt(12)

                        # Заголовки таблиц — жирные
                        run.bold = (row_idx == 0)
                        rPr = (run._element.get_or_add_rPr())
                        rFonts = rPr.rFonts
                        if rFonts is None:
                            rFonts = OxmlElement("w:rFonts")
                            rPr.append(rFonts)
                        rFonts.set(qn("w:eastAsia"),"Times New Roman")

    def create_list(self, doc: Document, items: list):
        """Создание маркированного списка"""
        for item in items:
            para = doc.add_paragraph(style="List Bullet")
            para_format = (para.paragraph_format)
            para_format.left_indent = Inches(0.3)
            para_format.first_line_indent = Inches(-0.3)
            run = para.add_run(item)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.rFonts
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:eastAsia"),"Times New Roman")

    def create_system_characteristics_table(self, doc: Document, data_dict: dict):
        """Создание таблицы характеристик системы"""
        if not self.has_system_characteristics(data_dict):
            return

        def format_numeric(value):
            """Форматирование числа, всегда 2 знака после запятой"""
            if value is None or str(value).strip() == "":
                return "Не указано"
            try:
                number = float(str(value).replace(",", "."))
                return f"{number:.2f}"
            except (ValueError,TypeError):
                return str(value)

        data_rows = [
            ("Тип дренажной системы",str(data_dict.get("SYSTEM_TYPE","Не указано")),"-"),
            ("Материал труб",str(data_dict.get("PIPE_MATERIAL","Не указано")),"-"),
            ("Диаметр труб",format_numeric(data_dict.get("PIPE_DIAMETER")),"мм"),
            ("Глубина заложения",format_numeric(data_dict.get("PIPE_DEPTH")),"м"),
            ("Общая протяженность",format_numeric(data_dict.get("PIPE_LENGTH")),"м"),
            ("Год ввода в эксплуатацию",format_numeric(data_dict.get("PIPE_INSTALL_YEAR")),"-"),
            ("Количество колодцев",format_numeric(data_dict.get("MANHOLE_COUNT")),"шт"),
        ]

        table = doc.add_table(rows=len(data_rows) + 1, cols=3)
        table.autofit = False
        widths = [Inches(2.3), Inches(1.9), Inches(1.8)]

        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = width

        headers = ["Параметр", "Значение", "Единицы измерения"]

        for i, header_text in enumerate(headers):
            para = (table.rows[0].cells[i].paragraphs[0])
            para.text = header_text
            para.alignment = (WD_ALIGN_PARAGRAPH.CENTER)

            for run in para.runs:
                run.font.name = ("Times New Roman")
                run.font.size = Pt(12)
                run.bold = True
                rPr = (run._element.get_or_add_rPr())
                rFonts = rPr.rFonts
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rPr.append(rFonts)
                rFonts.set(qn("w:eastAsia"),"Times New Roman")

        for i, row_data in enumerate(data_rows,start=1):
            cells = table.rows[i].cells
            for j, value in enumerate(row_data):
                para = (cells[j].paragraphs[0])
                para.text = str(value)
                para.alignment = (WD_ALIGN_PARAGRAPH.CENTER)
                for run in para.runs:
                    run.font.name = ("Times New Roman")
                    run.font.size = Pt(12)
                    rPr = (run._element.get_or_add_rPr())
                    rFonts = rPr.rFonts
                    if rFonts is None:
                        rFonts = OxmlElement("w:rFonts")
                        rPr.append(rFonts)
                    rFonts.set(qn("w:eastAsia"),"Times New Roman")
        self.style_table(table)

    def has_system_characteristics(self, data_dict: dict) -> bool:
        """Проверка наличия данных для раздела характеристик системы"""
        def has_positive_number(value):
            if value is None:
                return False
            try:
                return float(str(value).replace(',', '.')) > 0
            except ValueError:
                return False

        text_values = [data_dict.get('SYSTEM_TYPE'), data_dict.get('PIPE_MATERIAL')]
        if any(str(v).strip() for v in text_values):
            return True

        numeric_values = [data_dict.get('PIPE_DIAMETER'), data_dict.get('PIPE_DEPTH'), data_dict.get('PIPE_LENGTH')]
        if any(has_positive_number(v) for v in numeric_values):
            return True

        if has_positive_number(data_dict.get('PIPE_INSTALL_YEAR')):
            return True

        if has_positive_number(data_dict.get('MANHOLE_COUNT')):
            return True
        
        return False

    def create_table_element(self, doc: Document, table_type: str, data_dict: dict) -> bool:
        """Создание таблиц для разных типов данных"""
        if table_type == "SYSTEM_CHARACTERISTICS":
            if self.has_system_characteristics(data_dict):
                self.create_system_characteristics_table(doc, data_dict)
                return True
            return False

        elif table_type == "OBSERVATION_POINTS":
            points = []
            for item in data_dict.get("OBSERVATION_POINTS", []):
                if not isinstance(item, dict):
                    continue
                points.append((
                    str(item.get("observation_point") or item.get("name") or ""),
                    str(item.get("latitude") if item.get("latitude") is not None else item.get("lat") or ""),
                    str(item.get("longitude") if item.get("longitude") is not None else item.get("lon") or ""),
                    str(item.get("medium_type") or item.get("type") or ""),
                    str(item.get("description") or "")
                ))
            if points:
                self.create_observation_points_table(doc, points)

        elif table_type == "TEST_RESULTS":
            if data_dict.get("TEST_RESULTS"):
                self.create_test_results_table(doc, data_dict["TEST_RESULTS"])
                return True
            return False

        elif table_type == "OBSERVATION_DYNAMICS":
            if data_dict.get("OBSERVATION_DYNAMICS"):
                self.create_observation_dynamics_table(doc, data_dict["OBSERVATION_DYNAMICS"])
                return True
            return False

    def create_observation_points_table(self, doc: Document, points: list):
        """Таблица 2. Координаты точек наблюдения"""

        table = doc.add_table(rows=len(points) + 1, cols=6)
        headers = ["№", "Точка наблюдения", "Широта", "Долгота", "Тип среды", "Описание"]

        for i, header_text in enumerate(headers):
            table.rows[0].cells[i].text = header_text

        for i, point_data in enumerate(points, start=1):
            row = table.rows[i].cells
            row[0].text = str(i)
            row[1].text = str(point_data[0])
            row[2].text = str(point_data[1])
            row[3].text = str(point_data[2])
            row[4].text = str(point_data[3])
            row[5].text = str(point_data[4])

        self.style_table(table)

    def create_test_results_table(self, doc: Document, results: list):
        """Таблица 3. Результаты лабораторных анализов"""
        if not results:
            return

        table = doc.add_table(rows=len(results) + 1, cols=5)
        headers = ["Показатель", "Норматив", "Результат", "Единицы измерения", "Соответствие"]

        for i, header_text in enumerate(headers):
            table.rows[0].cells[i].text = header_text

        for i, result in enumerate(results, start=1):
            row = table.rows[i].cells
            indicator = str(result.get("indicator") or result.get("name") or "")
            
            # Норматив
            normative = (result.get("normative") or result.get("norm") or result.get("limit"))
            if not normative:
                normative = self.get_normative_range(indicator)

            # Результат
            result_value = (result.get("result") or result.get("value") or "")
            row[0].text = indicator
            row[1].text = str(normative)
            row[2].text = self.format_number(result_value)
            row[3].text = str(result.get("unit") or result.get("units") or result.get("measurement_unit") or "" )
            row[4].text = str(result.get("compliance") or "")

        self.style_table(table)

    def create_observation_dynamics_table(self, doc: Document, dynamics: list):
        """Таблица 4. Динамика наблюдений"""
        if not dynamics:
            return
        base_fields = [("pH", "pH"), ("iron", "Железо"), ("manganese", "Марганец"), ("nitrates", "Нитраты"), ("sulfates", "Сульфаты"),]
        used_fields = [("date", "Дата")]

        for field, label in base_fields:
            if any(self.has_value(item.get(field)) for item in dynamics):
                used_fields.append((field, label))

        # Если кроме даты ничего нет — не рисуем таблицу
        if len(used_fields) == 1:
            return

        table = doc.add_table(rows=len(dynamics) + 1, cols=len(used_fields))
        headers = [label for _, label in used_fields]

        # Заголовок
        for i, h in enumerate(headers):
            para = table.rows[0].cells[i].paragraphs[0]
            para.text = h
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

        # Данные
        for i, entry in enumerate(dynamics, start=1):
            cells = table.rows[i].cells
            for j, (field, _) in enumerate(used_fields):
                value = entry.get(field)
                if field == "date":
                    value = self.format_date_ddmmyyyy(value)
                else:
                    value = self.format_number(value)
                para = cells[j].paragraphs[0]
                para.text = str(value)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
        self.style_table(table)

    def read_section(self, file_path: str, data_dict: dict):
        """Чтение и обработка шаблонов разделов"""
        elements = []

        if not os.path.exists(file_path):
            return elements

        # Обработка специальных полей перед рендерингом
        processed_data = data_dict.copy()
        if isinstance(processed_data.get("DOCUMENTS_GOST"), list):
            processed_data["DOCUMENTS_GOST"] = "\n".join(processed_data["DOCUMENTS_GOST"])

        with open(file_path, "r", encoding="utf-8") as f:
            template = Template(f.read())

        rendered = template.render(**processed_data)

        buffer_list = []
        in_list = False

        for line in rendered.splitlines():
            line = line.strip()

            if not line:
                if buffer_list and in_list:
                    elements.append(("LIST", buffer_list.copy()))
                    buffer_list, in_list = [], False
                continue

            if line.startswith("TITLE:"):
                elements.append(("TITLE", line.replace("TITLE:", "").strip()))

            elif line.startswith("PARA:"):
                text = line.replace("PARA:", "").strip()
                text = self.remove_graph_desc_prefix(text)
                elements.append(("PARA", text))

            elif line.startswith("RIGHT_PARA:"):
                elements.append(("RIGHT_PARA", line.replace("RIGHT_PARA:", "").strip()))

            elif line.startswith("CENTER_PARA:"):
                elements.append(("CENTER_PARA", line.replace("CENTER_PARA:", "").strip()))

            elif line.startswith("TABLE:"):
                table_type = line.replace("TABLE:", "").strip()
                elements.append(("TABLE", table_type))

            elif line.startswith("GRAPH:"):
                graph_type = line.replace("GRAPH:", "").strip()
                elements.append(("GRAPH", graph_type))

            elif line.startswith("LIST:"):

                if buffer_list and in_list:
                    elements.append(("LIST", buffer_list.copy()))

                in_list = True
                buffer_list = []

            elif in_list:
                buffer_list.append(line.lstrip('•-*').strip())

            else:
                line = self.remove_graph_desc_prefix(line)
                elements.append(("PARA", line))

        if buffer_list and in_list:
            elements.append(("LIST", buffer_list))

        return elements

    def get_graph_image_stream(self, graph_type, data_dict):
        """Генерация графика в BytesIO для DOCX"""
        try:
            results = data_dict.get("TEST_RESULTS", [])
            dynamics = data_dict.get("OBSERVATION_DYNAMICS", [])
            image_stream = None

            if graph_type == "TEST_RESULTS":
                image_stream = comparison_bar_chart_docx(results)

            elif "OBSERVATION_DYNAMICS" in graph_type:
                metric = graph_type.replace("OBSERVATION_DYNAMICS:", "").strip().lower()
                image_stream = concentration_dynamics_lineplot_docx(results, dynamics, metric)

            if image_stream is None:
                logger.debug(f"График {graph_type} не создан (нет данных)")
                return None

            image_stream.seek(0)
            return image_stream

        except Exception as ex:
            logger.error(f"Ошибка генерации графика {graph_type}: {ex}", exc_info=True)
            return None
    
    def add_section_elements(self, doc: Document, elements: list, data_dict: dict):
        """Добавление элементов раздела"""

        def apply_font(run, size=12, bold=False):
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)
            run.bold = bold
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.rFonts
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), "Times New Roman")
        table_created = False
        for element_type, content in elements:
            if element_type == "TABLE":
                before_count = len(doc.tables)
                self.create_table_element(doc, content, data_dict)
                after_count = len(doc.tables)
                table_created = after_count > before_count
                continue
            if element_type == "RIGHT_PARA":
                if not table_created:
                    continue
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = para.add_run(content)
                apply_font(run)
                continue
            if element_type == "TITLE":
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(12)
                run = para.add_run(content)
                apply_font(run, size=14, bold=True)
                continue
            if element_type == "PARA":
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.first_line_indent = self.FIRST_LINE_INDENT
                para.paragraph_format.line_spacing = 1.5
                run = para.add_run(content)
                apply_font(run)
                continue
            if element_type == "CENTER_PARA":
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(content)
                apply_font(run)
                continue
            if element_type == "LIST":
                self.create_list(doc, content)
                continue
            if element_type == "GRAPH":
                try:
                    image_stream = self.get_graph_image_stream(content, data_dict)
                    if image_stream:
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run()
                        image_stream.seek(0)
                        run.add_picture(image_stream, width=Inches(6.5))
                        logger.debug(f"График {content} успешно вставлен в DOCX")
                    else:
                        logger.warning(f"Не удалось создать график {content} (нет данных или ошибка)")
                except Exception as ex:
                    logger.error(f"Критическая ошибка при вставке графика {content}: {ex}", exc_info=True)

    def generate(self, user_data: dict) -> bytes:
        """Основная функция генерации DOCX"""
        doc = Document()

        # Форматирование даты
        user_data = user_data.copy()
        user_data["REPORT_DATE"] = self.format_date_ddmmyyyy(user_data.get("REPORT_DATE"))

        # Установка полей
        self.set_margins(doc)

        # Титульная страница
        self.add_title_page(doc, user_data)

        # Разделы
        content_file = os.path.join(self.BASE_DIR, "report_module", "content.txt")

        if os.path.exists(content_file):
            with open(content_file, "r", encoding="utf-8") as f:
                sections = [line.strip() for line in f if line.strip()]

            for i, section in enumerate(sections, 1):
                section_file = os.path.join(self.BASE_DIR, "report_module", f"section_{i}.txt")
                elements = self.read_section(section_file, user_data)
                self.add_section_elements(doc, elements, user_data)
                if i < len(sections):
                    doc.add_page_break()

        # Сохранение в BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
